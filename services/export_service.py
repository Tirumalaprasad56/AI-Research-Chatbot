import os
import uuid

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer
)

from reportlab.lib.styles import getSampleStyleSheet


from docx import Document




# =====================================
# EXPORT FOLDER
# =====================================

EXPORT_FOLDER = "exports"


os.makedirs(
    EXPORT_FOLDER,
    exist_ok=True
)





# =====================================
# CLEAN MARKDOWN
# =====================================


def clean_heading(text):

    return text.replace(
        "#",
        ""
    ).strip()





# =====================================
# CREATE PDF
# =====================================


def create_pdf(report):


    try:


        filename = os.path.join(

            EXPORT_FOLDER,

            f"Research_Report_{uuid.uuid4().hex[:6]}.pdf"

        )



        doc = SimpleDocTemplate(
            filename
        )



        styles = getSampleStyleSheet()



        story=[]



        for line in report.split("\n"):


            line=line.strip()



            if not line:

                story.append(
                    Spacer(1,12)
                )

                continue



            line=clean_heading(line)



            story.append(

                Paragraph(

                    line,

                    styles["BodyText"]

                )

            )



            story.append(
                Spacer(1,8)
            )



        doc.build(
            story
        )



        return filename



    except Exception as e:


        raise Exception(

            f"PDF export failed: {str(e)}"

        )







# =====================================
# CREATE WORD DOCUMENT
# =====================================


def create_docx(report):


    try:



        filename=os.path.join(

            EXPORT_FOLDER,

            f"Research_Report_{uuid.uuid4().hex[:6]}.docx"

        )



        document=Document()



        document.add_heading(

            "AI Research Report",

            level=1

        )




        for line in report.split("\n"):



            line=line.strip()



            if not line:

                continue




            if line.startswith("#"):


                document.add_heading(

                    clean_heading(line),

                    level=2

                )


            else:


                document.add_paragraph(

                    line

                )




        document.save(

            filename

        )



        return filename




    except Exception as e:



        raise Exception(

            f"DOCX export failed: {str(e)}"

        )