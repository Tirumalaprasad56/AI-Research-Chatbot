import os

from PyPDF2 import PdfReader

from docx import Document



# =====================================
# CLEAN TEXT
# =====================================


def clean_text(text):

    if not text:

        return ""


    text = text.replace(
        "\x00",
        ""
    )


    lines = [

        line.strip()

        for line in text.split("\n")

        if line.strip()

    ]


    return "\n".join(lines)





# =====================================
# READ PDF
# =====================================


def read_pdf(filepath):


    try:


        text = ""


        reader = PdfReader(
            filepath
        )


        for page_number, page in enumerate(reader.pages):


            content = page.extract_text()



            if content:

                text += (
                    content
                    +
                    "\n"
                )



        return clean_text(text)



    except Exception as e:


        raise Exception(
            f"PDF reading error: {str(e)}"
        )





# =====================================
# READ DOCX
# =====================================


def read_docx(filepath):


    try:


        text = ""


        doc = Document(
            filepath
        )


        # Read paragraphs

        for paragraph in doc.paragraphs:


            if paragraph.text.strip():


                text += (

                    paragraph.text

                    +

                    "\n"

                )



        # Read tables

        for table in doc.tables:


            for row in table.rows:


                row_text = []


                for cell in row.cells:


                    row_text.append(
                        cell.text
                    )


                text += (

                    " | ".join(row_text)

                    +

                    "\n"

                )



        return clean_text(text)




    except Exception as e:


        raise Exception(
            f"DOCX reading error: {str(e)}"
        )





# =====================================
# READ TEXT FILE
# =====================================


def read_txt(filepath):


    try:


        with open(

            filepath,

            "r",

            encoding="utf-8",

            errors="ignore"

        ) as file:


            text = file.read()



        return clean_text(text)




    except Exception as e:


        raise Exception(
            f"TXT reading error: {str(e)}"
        )






# =====================================
# DOCUMENT ROUTER
# =====================================


def read_document(filepath):


    if not os.path.exists(filepath):


        raise Exception(
            "File does not exist"
        )



    extension = os.path.splitext(
        filepath
    )[1].lower()



    if extension == ".pdf":


        return read_pdf(filepath)



    elif extension == ".docx":


        return read_docx(filepath)



    elif extension == ".txt":


        return read_txt(filepath)



    else:


        raise Exception(
            "Unsupported document format"
        )